#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
report_engine.py — Multi-format report generator for QCD bridge results.

Outputs in 7 formats: TXT, CSV, MD, PDF, HTML, DOCX, JSON.
Structure: [RESULTS] first, then [LOGS] (full execution log).

Author: Ishak Khamzatovich Isaev (ORCID 0009-0003-7299-0701)
"""
from __future__ import annotations

import csv
import json
import logging
import os
import time
from dataclasses import asdict, is_dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import numpy as np

# Optional heavy imports for PDF/DOCX (lazy-loaded, but provide module-level fallbacks)
try:
    from reportlab.lib.units import cm as RL_CM
except ImportError:
    RL_CM = 28.3464567  # 1 cm in points fallback

logger = logging.getLogger("qcd_bridge.reports")


def _to_serializable(obj: Any) -> Any:
    """Recursively convert numpy/dataclass objects to JSON-serializable types."""
    if isinstance(obj, dict):
        return {k: _to_serializable(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_to_serializable(v) for v in obj]
    if isinstance(obj, (np.integer,)):
        return int(obj)
    if isinstance(obj, (np.floating,)):
        return float(obj)
    if isinstance(obj, np.ndarray):
        return obj.tolist()
    if isinstance(obj, (np.bool_,)):
        return bool(obj)
    if is_dataclass(obj):
        return _to_serializable(asdict(obj))
    if isinstance(obj, datetime):
        return obj.isoformat()
    return obj


class ReportEngine:
    """Generate reports in 7 formats with consistent structure: RESULTS then LOGS."""

    AUTHOR = "Ishak Khamzatovich Isaev"
    ORCID = "0009-0003-7299-0701"

    def __init__(self, output_dir: str = "reports", language: str = "en"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.language = language  # en | ru
        # i18n strings
        self.i18n = {
            "en": {
                "title": "Choptuik-QCD Bridge Verification Report",
                "results_section": "RESULTS",
                "logs_section": "EXECUTION LOG",
                "metadata_section": "METADATA",
                "generated": "Generated",
                "author": "Author",
                "orcid": "ORCID",
                "elapsed": "Elapsed (s)",
                "sections_run": "Sections executed",
                "config": "Configuration",
                "mode": "Mode",
                "summary": "Summary",
                "no_data": "No data",
            },
            "ru": {
                "title": "Отчёт о верификации моста Чоптуика–КХД",
                "results_section": "РЕЗУЛЬТАТЫ",
                "logs_section": "ЛОГ ВЫПОЛНЕНИЯ",
                "metadata_section": "МЕТАДАННЫЕ",
                "generated": "Создано",
                "author": "Автор",
                "orcid": "ORCID",
                "elapsed": "Время (с)",
                "sections_run": "Выполненные разделы",
                "config": "Конфигурация",
                "mode": "Режим",
                "summary": "Сводка",
                "no_data": "Нет данных",
            },
        }

    @property
    def tr(self):
        return self.i18n.get(self.language, self.i18n["en"])

    def generate(self, result, formats: Optional[List[str]] = None) -> Dict[str, str]:
        """Generate reports in all requested formats.

        Args:
            result: QCDBridgeResult dataclass
            formats: list of formats from ['txt','csv','md','pdf','html','docx','json']
                     If None, generates all 7.

        Returns:
            Dict mapping format -> file path
        """
        if formats is None:
            formats = ["txt", "csv", "md", "pdf", "html", "docx", "json"]

        # Normalize to JSON-safe structure
        data = _to_serializable(result)
        paths: Dict[str, str] = {}
        for fmt in formats:
            try:
                method = getattr(self, f"_write_{fmt}")
                path = method(data)
                paths[fmt] = path
                logger.info(f"Generated {fmt.upper()} report: {path}")
            except Exception as e:
                logger.error(f"Failed to generate {fmt} report: {e}")
                paths[fmt] = f"ERROR: {e}"
        return paths

    # ── JSON ─────────────────────────────────────────────────────────────
    def _write_json(self, data: Dict) -> str:
        path = self.output_dir / "report.json"
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False, default=str)
        return str(path)

    # ── TXT ──────────────────────────────────────────────────────────────
    def _write_txt(self, data: Dict) -> str:
        path = self.output_dir / "report.txt"
        lines = []
        lines.append("=" * 78)
        lines.append(f"  {self.tr['title']}")
        lines.append("=" * 78)
        lines.append(f"{self.tr['author']}: {self.AUTHOR}")
        lines.append(f"{self.tr['orcid']}: {self.ORCID}")
        lines.append(f"{self.tr['generated']}: {data.get('timestamp', '')}")
        lines.append(f"{self.tr['elapsed']}: {data.get('elapsed_s', 0):.3f}")
        lines.append(f"{self.tr['sections_run']}: {data.get('sections_run', [])}")
        lines.append(f"{self.tr['mode']}: {data.get('config', {}).get('mode', 'unknown')}")
        lines.append("")
        lines.append("=" * 78)
        lines.append(f"  {self.tr['results_section']}")
        lines.append("=" * 78)
        lines.append("")
        results = data.get("results", {})
        for section_key, section_data in results.items():
            lines.append("-" * 60)
            lines.append(f"  {section_key}")
            lines.append("-" * 60)
            self._dump_dict_txt(section_data, lines, indent=2)
            lines.append("")
        lines.append("")
        lines.append("=" * 78)
        lines.append(f"  {self.tr['logs_section']}")
        lines.append("=" * 78)
        for log in data.get("logs", []):
            lines.append(log)
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return str(path)

    def _dump_dict_txt(self, d: Any, lines: List[str], indent: int = 0):
        pad = " " * indent
        if isinstance(d, dict):
            for k, v in d.items():
                if isinstance(v, (dict, list)) and v:
                    lines.append(f"{pad}{k}:")
                    self._dump_dict_txt(v, lines, indent + 2)
                else:
                    val_str = self._fmt_val(v)
                    lines.append(f"{pad}{k}: {val_str}")
        elif isinstance(d, list):
            for i, item in enumerate(d):
                if isinstance(item, (dict, list)) and item:
                    lines.append(f"{pad}[{i}]:")
                    self._dump_dict_txt(item, lines, indent + 2)
                else:
                    lines.append(f"{pad}[{i}]: {self._fmt_val(item)}")
        else:
            lines.append(f"{pad}{self._fmt_val(d)}")

    def _fmt_val(self, v: Any) -> str:
        if isinstance(v, float):
            if abs(v) < 1e-3 and v != 0:
                return f"{v:.6e}"
            return f"{v:.6f}"
        if isinstance(v, list) and len(v) > 10:
            return f"[{len(v)} items, first 5: {v[:5]}...]"
        return str(v)

    # ── CSV ──────────────────────────────────────────────────────────────
    def _write_csv(self, data: Dict) -> str:
        path = self.output_dir / "report.csv"
        results = data.get("results", {})
        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["section", "key", "value", "type"])
            for section_key, section_data in results.items():
                self._dump_csv_rows(section_key, section_data, writer, prefix="")
            # Append log rows
            writer.writerow([])
            writer.writerow(["LOG", "timestamp", "message", ""])
            for log in data.get("logs", []):
                # log format: "[TS] message"
                writer.writerow(["LOG", "", log, "str"])
        return str(path)

    def _dump_csv_rows(self, section: str, d: Any, writer, prefix: str = ""):
        if isinstance(d, dict):
            for k, v in d.items():
                key = f"{prefix}.{k}" if prefix else k
                if isinstance(v, dict):
                    self._dump_csv_rows(section, v, writer, prefix=key)
                elif isinstance(v, list):
                    if v and isinstance(v[0], dict):
                        for i, item in enumerate(v):
                            self._dump_csv_rows(section, item, writer, prefix=f"{key}[{i}]")
                    else:
                        writer.writerow([section, key, self._fmt_val(v), "list"])
                else:
                    writer.writerow([section, key, self._fmt_val(v), type(v).__name__])
        elif isinstance(d, list):
            for i, item in enumerate(d):
                self._dump_csv_rows(section, item, writer, prefix=f"{prefix}[{i}]")

    # ── Markdown ─────────────────────────────────────────────────────────
    def _write_md(self, data: Dict) -> str:
        path = self.output_dir / "report.md"
        lines = []
        lines.append(f"# {self.tr['title']}")
        lines.append("")
        lines.append(f"**{self.tr['author']}**: {self.AUTHOR}  ")
        lines.append(f"**{self.tr['orcid']}**: {self.ORCID}  ")
        lines.append(f"**{self.tr['generated']}**: {data.get('timestamp', '')}  ")
        lines.append(f"**{self.tr['elapsed']}**: {data.get('elapsed_s', 0):.3f}  ")
        lines.append(f"**{self.tr['mode']}**: `{data.get('config', {}).get('mode', '')}`  ")
        lines.append(f"**{self.tr['sections_run']}**: {data.get('sections_run', [])}  ")
        lines.append("")
        lines.append("---")
        lines.append("")
        lines.append(f"## {self.tr['results_section']}")
        lines.append("")
        results = data.get("results", {})
        for section_key, section_data in results.items():
            lines.append(f"### {section_key}")
            lines.append("")
            self._dump_md(section_data, lines, level=4)
            lines.append("")
        lines.append("---")
        lines.append("")
        lines.append(f"## {self.tr['logs_section']}")
        lines.append("")
        lines.append("```")
        for log in data.get("logs", []):
            lines.append(log)
        lines.append("```")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return str(path)

    def _dump_md(self, d: Any, lines: List[str], level: int = 4):
        header = "#" * level
        if isinstance(d, dict):
            for k, v in d.items():
                if isinstance(v, dict) and v:
                    lines.append(f"{header} {k}")
                    lines.append("")
                    self._dump_md(v, lines, level + 1)
                elif isinstance(v, list) and v and isinstance(v[0], dict):
                    lines.append(f"{header} {k}")
                    lines.append("")
                    for i, item in enumerate(v):
                        lines.append(f"{header}# Item {i}")
                        lines.append("")
                        self._dump_md(item, lines, level + 1)
                else:
                    val_str = self._fmt_val(v)
                    lines.append(f"- **{k}**: `{val_str}`")
            lines.append("")
        elif isinstance(d, list):
            for i, item in enumerate(d):
                lines.append(f"{header} Item {i}: `{self._fmt_val(item)}`")
            lines.append("")

    # ── HTML ─────────────────────────────────────────────────────────────
    def _write_html(self, data: Dict) -> str:
        path = self.output_dir / "report.html"
        results = data.get("results", {})
        html = []
        html.append("<!DOCTYPE html>")
        html.append("<html lang='" + self.language + "'><head><meta charset='UTF-8'>")
        html.append(f"<title>{self.tr['title']}</title>")
        html.append("<style>")
        html.append("body{font-family:'Segoe UI',Roboto,Arial,sans-serif;max-width:1100px;margin:2em auto;padding:1em;color:#1F2937;background:#FFFFFF;}")
        html.append("h1{color:#101820;border-bottom:3px solid #4C6EF5;padding-bottom:0.3em;}")
        html.append("h2{color:#243447;margin-top:2em;}")
        html.append("h3{color:#3AAFA9;}")
        html.append("table{border-collapse:collapse;width:100%;margin:1em 0;}")
        html.append("th,td{border:1px solid #E5E7EB;padding:0.5em 0.8em;text-align:left;font-size:0.92em;}")
        html.append("th{background:#F1F5F9;color:#243447;}")
        html.append("tr:nth-child(even){background:#F8FAFC;}")
        html.append(".meta{background:#F1F5F9;padding:1em 1.5em;border-radius:6px;margin:1em 0;}")
        html.append(".meta span{display:inline-block;margin-right:1.5em;}")
        html.append(".log{background:#0F172A;color:#F1F5F9;padding:1em;border-radius:6px;font-family:'Fira Code','Consolas',monospace;font-size:0.85em;overflow-x:auto;white-space:pre-wrap;}")
        html.append(".value{font-family:'Fira Code','Consolas',monospace;color:#37352F;}")
        html.append("</style></head><body>")
        html.append(f"<h1>{self.tr['title']}</h1>")
        html.append("<div class='meta'>")
        html.append(f"<span><b>{self.tr['author']}:</b> {self.AUTHOR}</span>")
        html.append(f"<span><b>{self.tr['orcid']}:</b> <a href='https://orcid.org/{self.ORCID}'>{self.ORCID}</a></span>")
        html.append(f"<span><b>{self.tr['generated']}:</b> {data.get('timestamp','')}</span>")
        html.append(f"<span><b>{self.tr['elapsed']}:</b> {data.get('elapsed_s',0):.3f}</span>")
        html.append(f"<span><b>{self.tr['mode']}:</b> <code>{data.get('config',{}).get('mode','')}</code></span>")
        html.append("</div>")
        html.append(f"<h2>{self.tr['results_section']}</h2>")
        for section_key, section_data in results.items():
            html.append(f"<h3>{section_key}</h3>")
            html.append(self._dump_html(section_data))
        html.append(f"<h2>{self.tr['logs_section']}</h2>")
        html.append("<div class='log'>")
        for log in data.get("logs", []):
            # escape HTML
            safe = log.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html.append(safe + "\n")
        html.append("</div>")
        html.append("</body></html>")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(html))
        return str(path)

    def _dump_html(self, d: Any) -> str:
        if isinstance(d, dict):
            rows = []
            for k, v in d.items():
                if isinstance(v, dict) and v:
                    rows.append(f"<tr><td><b>{k}</b></td><td>{self._dump_html(v)}</td></tr>")
                elif isinstance(v, list) and v and isinstance(v[0], dict):
                    inner = "".join(
                        f"<tr><td><b>[{i}]</b></td><td>{self._dump_html(item)}</td></tr>"
                        for i, item in enumerate(v)
                    )
                    rows.append(f"<tr><td><b>{k}</b></td><td><table>{inner}</table></td></tr>")
                else:
                    val_str = self._fmt_val(v).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    rows.append(f"<tr><td><b>{k}</b></td><td><span class='value'>{val_str}</span></td></tr>")
            return "<table>" + "".join(rows) + "</table>"
        if isinstance(d, list):
            items = "".join(f"<li><span class='value'>{self._fmt_val(x)}</span></li>" for x in d)
            return f"<ul>{items}</ul>"
        return f"<span class='value'>{self._fmt_val(d)}</span>"

    # ── PDF ──────────────────────────────────────────────────────────────
    def _write_pdf(self, data: Dict) -> str:
        path = self.output_dir / "report.pdf"
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Preformatted
            )
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_LEFT
            cm = RL_CM
        except ImportError:
            # Fallback: write a minimal PDF stub
            with open(path, "wb") as f:
                f.write(b"%PDF-1.4\n% ReportLab not installed; stub only\n")
            return str(path)

        doc = SimpleDocTemplate(
            str(path), pagesize=A4,
            leftMargin=2 * cm, rightMargin=2 * cm,
            topMargin=2 * cm, bottomMargin=2 * cm,
            title=self.tr["title"], author=self.AUTHOR,
        )
        styles = getSampleStyleSheet()
        h1 = styles["Heading1"]
        h2 = styles["Heading2"]
        h3 = styles["Heading3"]
        body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10, leading=14, alignment=TA_LEFT)
        mono = ParagraphStyle("Mono", parent=styles["Code"], fontSize=8, leading=10)
        story = []
        story.append(Paragraph(self.tr["title"], h1))
        story.append(Spacer(1, 6))
        meta_rows = [
            [self.tr["author"], self.AUTHOR],
            [self.tr["orcid"], self.ORCID],
            [self.tr["generated"], data.get("timestamp", "")],
            [self.tr["elapsed"], f"{data.get('elapsed_s', 0):.3f} s"],
            [self.tr["mode"], data.get("config", {}).get("mode", "")],
            [self.tr["sections_run"], ", ".join(str(s) for s in data.get("sections_run", []))],
        ]
        t = Table(meta_rows, colWidths=[4 * cm, 12 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F1F5F9")),
            ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#243447")),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E5E7EB")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))
        story.append(Paragraph(self.tr["results_section"], h2))
        results = data.get("results", {})
        for section_key, section_data in results.items():
            story.append(Paragraph(section_key, h3))
            self._dump_pdf(section_data, story, styles, body, depth=0)
            story.append(Spacer(1, 6))
        story.append(PageBreak())
        story.append(Paragraph(self.tr["logs_section"], h2))
        log_text = "\n".join(data.get("logs", []))
        story.append(Preformatted(log_text, mono))
        doc.build(story)
        return str(path)

    def _dump_pdf(self, d: Any, story, styles, body_style, depth: int = 0):
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        cm = RL_CM
        if isinstance(d, dict):
            rows = []
            for k, v in d.items():
                if isinstance(v, dict) and v:
                    story.append(Paragraph(f"<b>{k}</b>", body_style))
                    self._dump_pdf(v, story, styles, body_style, depth + 1)
                elif isinstance(v, list) and v and isinstance(v[0], dict):
                    story.append(Paragraph(f"<b>{k}</b>", body_style))
                    for i, item in enumerate(v):
                        story.append(Paragraph(f"<i>[{i}]</i>", body_style))
                        self._dump_pdf(item, story, styles, body_style, depth + 1)
                else:
                    val_str = self._fmt_val(v).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    rows.append([str(k), val_str])
            if rows:
                t = Table(rows, colWidths=[5 * cm, 11 * cm])
                t.setStyle(TableStyle([
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E5E7EB")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 2),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]))
                story.append(t)
                story.append(Spacer(1, 4))

    # ── DOCX ─────────────────────────────────────────────────────────────
    def _write_docx(self, data: Dict) -> str:
        path = self.output_dir / "report.docx"
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            with open(path, "wb") as f:
                f.write(b"python-docx not installed; install with: pip install python-docx\n")
            return str(path)

        doc = Document()
        # Title
        title = doc.add_heading(self.tr["title"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Meta
        meta = doc.add_table(rows=6, cols=2)
        meta.style = "Light Grid Accent 1"
        meta_rows = [
            (self.tr["author"], self.AUTHOR),
            (self.tr["orcid"], self.ORCID),
            (self.tr["generated"], data.get("timestamp", "")),
            (self.tr["elapsed"], f"{data.get('elapsed_s', 0):.3f} s"),
            (self.tr["mode"], data.get("config", {}).get("mode", "")),
            (self.tr["sections_run"], ", ".join(str(s) for s in data.get("sections_run", []))),
        ]
        for i, (k, v) in enumerate(meta_rows):
            meta.rows[i].cells[0].text = k
            meta.rows[i].cells[1].text = str(v)
        doc.add_paragraph()
        # Results
        doc.add_heading(self.tr["results_section"], level=1)
        results = data.get("results", {})
        for section_key, section_data in results.items():
            doc.add_heading(section_key, level=2)
            self._dump_docx(section_data, doc, depth=0)
        # Logs
        doc.add_page_break()
        doc.add_heading(self.tr["logs_section"], level=1)
        log_para = doc.add_paragraph()
        log_para.style = doc.styles["No Spacing"]
        run = log_para.add_run("\n".join(data.get("logs", [])))
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        doc.save(str(path))
        return str(path)

    def _dump_docx(self, d: Any, doc, depth: int = 0):
        from docx.shared import Pt
        if isinstance(d, dict):
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Light List Accent 1"
            for k, v in d.items():
                if isinstance(v, dict) and v:
                    row = tbl.add_row().cells
                    row[0].text = str(k)
                    row[1].text = ""
                    self._dump_docx(v, doc, depth + 1)
                elif isinstance(v, list) and v and isinstance(v[0], dict):
                    doc.add_paragraph(str(k)).bold = True
                    for i, item in enumerate(v):
                        doc.add_paragraph(f"[{i}]").italic = True
                        self._dump_docx(item, doc, depth + 1)
                else:
                    row = tbl.add_row().cells
                    row[0].text = str(k)
                    row[1].text = self._fmt_val(v)
                    for cell in row:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(9)
        elif isinstance(d, list):
            for i, item in enumerate(d):
                p = doc.add_paragraph()
                p.add_run(f"[{i}] ").bold = True
                p.add_run(self._fmt_val(item))


if __name__ == "__main__":
    # Smoke test
    import sys
    sys.path.insert(0, str(Path(__file__).parent))
    from qcd_bridge_engine import QCDBridgeConfig, run_all

    logging.basicConfig(level=logging.INFO)
    cfg = QCDBridgeConfig(mode="verify_all", sections=[1, 5, 6, 7, 8, 9])  # quick subset
    result = run_all(cfg)
    engine = ReportEngine(output_dir="/home/z/my-project/choptuik_ac_bc/qcd_bridge/reports", language="en")
    paths = engine.generate(result, formats=["json", "txt", "md", "html", "csv"])
    for fmt, p in paths.items():
        print(f"  {fmt}: {p}")
