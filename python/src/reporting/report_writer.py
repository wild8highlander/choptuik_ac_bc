"""Multi-format report generation: DOCX, PDF, TXT, MD, CSV, HTML, JSON.

Each report contains:
1. Results section - computed constants, deviations, comparison tables
2. Execution log - complete timestamped log of all computations

All reports are saved as separate files in the output directory.
"""

from __future__ import annotations
import json
import csv
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime
import logging

logger = logging.getLogger(__name__)


class ReportWriter:
    """Generate reports in 7 formats with execution logs appended."""

    def __init__(self, output_dir: str = "output/reports",
                 formats: Optional[List[str]] = None):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.formats = formats or ["docx", "pdf", "txt", "md", "csv", "html", "json"]
        logger.info(f"ReportWriter: output={self.output_dir}, formats={self.formats}")

    def generate_all(self, results: Dict, logs: str,
                     title: str = "Choptyuk Spinor Corrections - Verification Report") -> Dict[str, str]:
        """Generate reports in all configured formats."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        paths = {}
        for fmt in self.formats:
            try:
                if fmt == "json":
                    paths[fmt] = self._write_json(results, logs, timestamp)
                elif fmt == "txt":
                    paths[fmt] = self._write_txt(results, logs, title, timestamp)
                elif fmt == "md":
                    paths[fmt] = self._write_md(results, logs, title, timestamp)
                elif fmt == "csv":
                    paths[fmt] = self._write_csv(results, logs, timestamp)
                elif fmt == "html":
                    paths[fmt] = self._write_html(results, logs, title, timestamp)
                elif fmt == "docx":
                    paths[fmt] = self._write_docx(results, logs, title, timestamp)
                elif fmt == "pdf":
                    paths[fmt] = self._write_pdf(results, logs, title, timestamp)
                else:
                    logger.warning(f"Unknown format: {fmt}")
            except Exception as e:
                logger.error(f"Failed to generate {fmt} report: {e}")
        return paths

    def _write_json(self, results: Dict, logs: str, timestamp: str) -> str:
        data = {
            "report_type": "choptyuk_verification",
            "timestamp": timestamp,
            "results": results,
            "execution_log": logs,
        }
        path = self.output_dir / f"report_{timestamp}.json"
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2, default=str)
        return str(path)

    def _write_txt(self, results: Dict, logs: str, title: str, timestamp: str) -> str:
        path = self.output_dir / f"report_{timestamp}.txt"
        with open(path, 'w', encoding='utf-8') as f:
            f.write(f"{'='*60}\n{title}\nGenerated: {timestamp}\n{'='*60}\n\n")
            f.write(self._format_dict(results))
            f.write(f"\n{'='*60}\nEXECUTION LOG\n{'='*60}\n")
            f.write(logs)
        return str(path)

    def _format_dict(self, d: Dict, indent: int = 0) -> str:
        lines = []
        prefix = "  " * indent
        for k, v in d.items():
            if isinstance(v, dict):
                lines.append(f"{prefix}{k}:")
                lines.append(self._format_dict(v, indent + 1))
            elif isinstance(v, list):
                lines.append(f"{prefix}{k}: [{len(v)} items]")
            else:
                if isinstance(v, float):
                    lines.append(f"{prefix}{k}: {v:.6f}")
                else:
                    lines.append(f"{prefix}{k}: {v}")
        return "\n".join(lines)

    def _write_md(self, results: Dict, logs: str, title: str, timestamp: str) -> str:
        path = self.output_dir / f"report_{timestamp}.md"
        ch = results.get("choptyuk", {})
        curve = results.get("curve", {})
        with open(path, 'w', encoding='utf-8') as f:
            f.write(f"# {title}\n\n**Generated**: {timestamp}\n\n## Results\n\n")
            if curve:
                f.write("### Klein Curve\n\n")
                f.write(f"- Genus: {curve.get('genus')}\n")
                f.write(f"- K = {curve.get('K')}, R = {curve.get('R')}\n")
                f.write(f"- Area = {curve.get('area', 0):.4f}\n")
                f.write(f"- |PSL(2,7)| = {curve.get('psl_order')}\n\n")
            if ch:
                f.write("### Choptyuk Formula\n\n")
                f.write("| Constant | Value | Deviation |\n|---|---|---|\n")
                f.write(f"| Delta_bC | {ch.get('delta_bc',0):.6f} | {ch.get('deviation_bc_pct',0):.3f}% |\n")
                f.write(f"| Delta_Ch (base) | {ch.get('delta_ch_base',0):.6f} | {ch.get('deviation_ch_pct',0):.3f}% |\n")
                f.write(f"| Delta_Ch (full) | {ch.get('delta_ch_full',0):.6f} | {ch.get('deviation_full_pct',0):.3f}% |\n")
                f.write(f"| b_Ch | {ch.get('b_ch',0):.6f} | {ch.get('deviation_b_ch_pct',0):.3f}% |\n\n")
            f.write("## Execution Log\n\n```\n" + logs + "\n```\n")
        return str(path)

    def _write_csv(self, results: Dict, logs: str, timestamp: str) -> str:
        path = self.output_dir / f"report_{timestamp}.csv"
        ch = results.get("choptyuk", {})
        with open(path, 'w', newline='', encoding='utf-8') as f:
            w = csv.writer(f)
            w.writerow(["constant", "value", "deviation_pct"])
            for name, val_key, dev_key in [
                ("delta_bc", "delta_bc", "deviation_bc_pct"),
                ("delta_ch_base", "delta_ch_base", "deviation_ch_pct"),
                ("delta_ch_full", "delta_ch_full", "deviation_full_pct"),
                ("b_ch", "b_ch", "deviation_b_ch_pct"),
            ]:
                w.writerow([name, ch.get(val_key, ""), ch.get(dev_key, "")])
            w.writerow([])
            w.writerow(["--- Execution Log ---"])
            for line in logs.split("\n"):
                w.writerow([line])
        return str(path)

    def _write_html(self, results: Dict, logs: str, title: str, timestamp: str) -> str:
        path = self.output_dir / f"report_{timestamp}.html"
        ch = results.get("choptyuk", {})
        log_escaped = logs.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{title}</title>
<style>
body{{font-family:sans-serif;max-width:900px;margin:0 auto;padding:20px}}
h1{{color:#2c3e50}}h2{{color:#34495e;border-bottom:2px solid #3498db;padding-bottom:5px}}
table{{border-collapse:collapse;width:100%;margin:10px 0}}
th,td{{border:1px solid #bdc3c7;padding:8px 12px;text-align:left}}
th{{background:#3498db;color:white}}tr:nth-child(even){{background:#f2f2f2}}
pre{{background:#2c3e50;color:#ecf0f1;padding:15px;border-radius:5px;overflow-x:auto;font-size:12px}}
</style></head><body>
<h1>{title}</h1><p><b>Generated</b>: {timestamp}</p>
<h2>Results</h2>
<table><tr><th>Constant</th><th>Value</th><th>Deviation</th></tr>
<tr><td>Delta_bC</td><td>{ch.get('delta_bc',0):.6f}</td><td>{ch.get('deviation_bc_pct',0):.3f}%</td></tr>
<tr><td>Delta_Ch (base)</td><td>{ch.get('delta_ch_base',0):.6f}</td><td>{ch.get('deviation_ch_pct',0):.3f}%</td></tr>
<tr><td>Delta_Ch (full)</td><td>{ch.get('delta_ch_full',0):.6f}</td><td>{ch.get('deviation_full_pct',0):.3f}%</td></tr>
<tr><td>b_Ch</td><td>{ch.get('b_ch',0):.6f}</td><td>{ch.get('deviation_b_ch_pct',0):.3f}%</td></tr>
</table>
<h2>Execution Log</h2><pre>{log_escaped}</pre>
</body></html>"""
        with open(path, 'w', encoding='utf-8') as f:
            f.write(html)
        return str(path)

    def _write_docx(self, results: Dict, logs: str, title: str, timestamp: str) -> str:
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed, skipping DOCX")
            return ""
        path = self.output_dir / f"report_{timestamp}.docx"
        doc = Document()
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Generated: {timestamp}")
        ch = results.get("choptyuk", {})
        doc.add_heading("Choptyuk Formula Results", level=1)
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        headers = ["Constant", "Value", "Deviation"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for i, (name, val, dev) in enumerate([
            ("Delta_bC", ch.get("delta_bc"), ch.get("deviation_bc_pct")),
            ("Delta_Ch (base)", ch.get("delta_ch_base"), ch.get("deviation_ch_pct")),
            ("Delta_Ch (full)", ch.get("delta_ch_full"), ch.get("deviation_full_pct")),
            ("b_Ch", ch.get("b_ch"), ch.get("deviation_b_ch_pct")),
        ], 1):
            table.rows[i].cells[0].text = name
            table.rows[i].cells[1].text = f"{val:.6f}" if val else ""
            table.rows[i].cells[2].text = f"{dev:.3f}%" if dev else ""
        doc.add_heading("Execution Log", level=1)
        doc.add_paragraph(logs, style='No Spacing')
        doc.save(str(path))
        return str(path)

    def _write_pdf(self, results: Dict, logs: str, title: str, timestamp: str) -> str:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, Preformatted
            from reportlab.lib.styles import getSampleStyleSheet
        except ImportError:
            logger.warning("reportlab not installed, skipping PDF")
            return ""
        path = self.output_dir / f"report_{timestamp}.pdf"
        doc = SimpleDocTemplate(str(path), pagesize=letter)
        styles = getSampleStyleSheet()
        ch = results.get("choptyuk", {})
        elems = [
            Paragraph(title, styles['Title']),
            Paragraph(f"Generated: {timestamp}", styles['Normal']),
            Spacer(1, 0.2 * inch),
            Paragraph("Choptyuk Formula Results", styles['Heading1']),
            Table([
                ['Constant', 'Value', 'Deviation'],
                ['Delta_bC', f"{ch.get('delta_bc',0):.6f}", f"{ch.get('deviation_bc_pct',0):.3f}%"],
                ['Delta_Ch (base)', f"{ch.get('delta_ch_base',0):.6f}", f"{ch.get('deviation_ch_pct',0):.3f}%"],
                ['Delta_Ch (full)', f"{ch.get('delta_ch_full',0):.6f}", f"{ch.get('deviation_full_pct',0):.3f}%"],
                ['b_Ch', f"{ch.get('b_ch',0):.6f}", f"{ch.get('deviation_b_ch_pct',0):.3f}%"],
            ], colWidths=[2*inch, 1.5*inch, 1.5*inch]),
            Spacer(1, 0.3 * inch),
            Paragraph("Execution Log", styles['Heading1']),
        ]
        for line in logs.split("\n")[-50:]:
            elems.append(Paragraph(line, styles['Code']))
        doc.build(elems)
        return str(path)
